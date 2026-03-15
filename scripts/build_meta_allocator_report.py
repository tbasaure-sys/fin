from __future__ import annotations

import json
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_ROOT = PROJECT_ROOT / "output"
RESEARCH_ROOT = OUTPUT_ROOT / "research" / "latest"
POLICY_ROOT = OUTPUT_ROOT / "policy" / "latest"
DASHBOARD_ROOT = OUTPUT_ROOT / "dashboard" / "latest"
DOC_ROOT = OUTPUT_ROOT / "doc"
TMP_ROOT = PROJECT_ROOT / "tmp" / "docs"


def _load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8").replace("NaN", "null"))


def _pct(value: float | None, digits: int = 1) -> str:
    if value is None or pd.isna(value):
        return "-"
    return f"{value * 100:.{digits}f}%"


def _num(value: float | None, digits: int = 2) -> str:
    if value is None or pd.isna(value):
        return "-"
    return f"{value:.{digits}f}"


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _style(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    for style_name, size in [("Normal", 10.5), ("Title", 21), ("Heading 1", 15), ("Heading 2", 12)]:
        style = document.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)


def _bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def _table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        _shade(cell, "D8EAF6")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def _plot_growth(series: pd.DataFrame, out_path: Path) -> None:
    plt.figure(figsize=(9.4, 4.8))
    colors = {
        "benchmark_spy": "#15324f",
        "state_overlay": "#0f8a6d",
        "policy_overlay": "#c06b1e",
        "trend_following": "#7057b7",
    }
    labels = {
        "benchmark_spy": "SPY",
        "state_overlay": "Heuristic state overlay",
        "policy_overlay": "Policy overlay",
        "trend_following": "Trend following",
    }
    for column in ["benchmark_spy", "state_overlay", "policy_overlay", "trend_following"]:
        if column in series.columns:
            plt.plot(series.index, series[column], linewidth=2.0, color=colors[column], label=labels[column])
    plt.title("Growth of $1")
    plt.ylabel("Cumulative wealth")
    plt.grid(alpha=0.22)
    plt.legend(frameon=False)
    plt.tight_layout()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    plt.savefig(out_path, dpi=180)
    plt.close()


def _plot_drawdown(series: pd.DataFrame, out_path: Path) -> None:
    plt.figure(figsize=(9.4, 4.8))
    for column, color, label in [
        ("benchmark_spy", "#15324f", "SPY"),
        ("state_overlay", "#0f8a6d", "Heuristic"),
        ("policy_overlay", "#c06b1e", "Policy"),
        ("trend_following", "#7057b7", "Trend"),
    ]:
        if column in series.columns:
            plt.plot(series.index, series[column], linewidth=2.0, color=color, label=label)
    plt.title("Drawdown")
    plt.ylabel("Peak-to-trough")
    plt.grid(alpha=0.22)
    plt.legend(frameon=False)
    plt.tight_layout()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    plt.savefig(out_path, dpi=180)
    plt.close()


def _title_block(document: Document, dashboard: dict) -> None:
    title = document.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Meta Allocator Methodology Report").bold = True

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Updated system design, research results, discovery stack, and deployment structure").italic = True

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"As of {dashboard['generated_at']} | Live regime {dashboard['overview'].get('regime', '-')}")


def _executive_summary(document: Document, research: dict, policy: dict, dashboard: dict) -> None:
    document.add_heading("1. Executive Summary", level=1)
    document.add_paragraph(
        "The project is now centered on a cleaner operating question than pure stock picking: how much equity risk should be carried now, "
        "which hedge should protect that exposure, and where the opportunity terrain is most attractive if discretionary capital is deployed."
    )
    _bullet(document, f"SPY buy and hold: CAGR {_pct(research['benchmark_spy']['annual_return'])}, Sharpe {_num(research['benchmark_spy']['sharpe'])}, max drawdown {_pct(research['benchmark_spy']['max_drawdown'])}.")
    _bullet(document, f"Heuristic state overlay: CAGR {_pct(research['state_overlay']['annual_return'])}, Sharpe {_num(research['state_overlay']['sharpe'])}, max drawdown {_pct(research['state_overlay']['max_drawdown'])}.")
    _bullet(document, f"Policy overlay: CAGR {_pct(policy['policy_overlay']['annual_return'])}, Sharpe {_num(policy['policy_overlay']['sharpe'])}, max drawdown {_pct(policy['policy_overlay']['max_drawdown'])}.")
    _bullet(document, f"Current live stance: beta target {_pct(dashboard['overview'].get('beta_target'))} with hedge {dashboard['overview'].get('selected_hedge', '-')}.")
    _bullet(document, "The policy overlay remains useful as a risk-aware operating layer, but the current strongest simple benchmark in the stack is still trend following.")


def _plain_language(document: Document, dashboard: dict) -> None:
    document.add_heading("2. The Idea In Plain Language", level=1)
    document.add_paragraph(
        "A normal investor in SPY stays almost fully invested all the time. This system tries to be more selective without becoming a black box. "
        "It asks whether the market environment is sturdy or fragile, whether consensus looks healthy or brittle, and which protection asset is best if risk should be reduced."
    )
    _bullet(document, "It does not try to predict every stock every day.")
    _bullet(document, "It does try to decide how much market exposure is justified.")
    _bullet(document, "It also tracks where opportunity is improving across sectors, international markets, and discovery names.")
    _bullet(document, f"Today the system sees regime `{dashboard['overview'].get('regime', '-')}`, consensus fragility {_pct(dashboard['overview'].get('consensus_fragility_score'))}, and belief-capacity mismatch {_pct(dashboard['overview'].get('belief_capacity_misalignment'))}.")


def _system_methodology(document: Document) -> None:
    document.add_heading("3. Core Methodology", level=1)
    _table(
        document,
        ["Component", "Role", "What it contributes"],
        [
            ["State engine", "Sense fragility", "Crash probability, legitimacy risk, tail risk, macro, cross-asset structure"],
            ["Hedge engine", "Choose protection", "Ranks SHY, IEF, TLT, BIL, GLD, UUP by conditional usefulness"],
            ["Policy overlay", "Choose exposure size", "Learns beta buckets instead of trying to forecast exact returns"],
            ["Scenario synthesis", "Stabilize decisions", "Blends learned policy with Bayesian scenario weighting"],
            ["Opportunity maps", "Context, not command", "Shows favored sectors and international markets"],
            ["Statement intelligence", "Economic quality", "Adds accounting quality, valuation context, and cash confirmation"],
            ["Discovery engine", "Find new names", "Separates current holdings from actual discovery ideas"],
        ],
    )

    document.add_heading("4. New Additions In This Version", level=1)
    _bullet(document, "Discovery was separated from holdings. Current positions no longer dominate the screener by construction.")
    _bullet(document, "Best daily-screen discoveries now receive valuation and simulation enrichment instead of staying purely technical.")
    _bullet(document, "Consensus fragility and belief-capacity mismatch were added as behavioral state variables.")
    _bullet(document, "Owner elasticity was added to rank names where smaller proprietary capital can plausibly matter more than institutional scale.")
    _bullet(document, "The earnings-to-cash kernel remains research-only and is additive to statement intelligence, not part of allocator core logic.")


def _data_section(document: Document) -> None:
    document.add_heading("5. Data And Feature Stack", level=1)
    _bullet(document, "Historical equity universe and membership history from local `caria_publication` data.")
    _bullet(document, "Systemic state from `Fin_model`.")
    _bullet(document, "Holdings, valuation, simulation, and discovery artifacts from `portfolio_manager`.")
    _bullet(document, "Live market proxies from FMP with yfinance fallback.")
    _bullet(document, "Macro and rates from FRED.")
    _bullet(document, "Quarterly statements for cash-confirmation work from FMP with yfinance fallback.")


def _results_section(document: Document, research: dict, policy: dict, dashboard: dict, growth_chart: Path, drawdown_chart: Path) -> None:
    document.add_heading("6. Research Results", level=1)
    document.add_paragraph("Core comparison across the current stack:")
    _table(
        document,
        ["Strategy", "CAGR", "Sharpe", "Max Drawdown"],
        [
            ["SPY", _pct(research["benchmark_spy"]["annual_return"]), _num(research["benchmark_spy"]["sharpe"]), _pct(research["benchmark_spy"]["max_drawdown"])],
            ["Selection standalone", _pct(research["selection_standalone"]["annual_return"]), _num(research["selection_standalone"]["sharpe"]), _pct(research["selection_standalone"]["max_drawdown"])],
            ["Heuristic state overlay", _pct(research["state_overlay"]["annual_return"]), _num(research["state_overlay"]["sharpe"]), _pct(research["state_overlay"]["max_drawdown"])],
            ["Policy overlay", _pct(policy["policy_overlay"]["annual_return"]), _num(policy["policy_overlay"]["sharpe"]), _pct(policy["policy_overlay"]["max_drawdown"])],
            ["Trend following", _pct(policy["trend_following"]["annual_return"]), _num(policy["trend_following"]["sharpe"]), _pct(policy["trend_following"]["max_drawdown"])],
        ],
    )
    document.add_paragraph("Growth comparison:")
    document.add_picture(str(growth_chart), width=Inches(6.8))
    document.add_paragraph("Drawdown comparison:")
    document.add_picture(str(drawdown_chart), width=Inches(6.8))

    document.add_heading("7. Out-of-Sample And Regime Lens", level=1)
    _table(
        document,
        ["Block", "Start", "End", "CAGR", "Sharpe", "MaxDD", "Episodes"],
        [
            [
                str(idx),
                block.get("start", "-"),
                block.get("end", "-"),
                _pct(block.get("annual_return")),
                _num(block.get("sharpe")),
                _pct(block.get("max_drawdown")),
                ", ".join(block.get("overlapping_episodes", [])[:3]) or "-",
            ]
            for idx, block in enumerate(research.get("oos_blocks", []), start=1)
        ],
    )
    regime_rows = research.get("regime_performance", [])
    if regime_rows:
        _table(
            document,
            ["Regime", "Obs", "CAGR", "Sharpe", "MaxDD"],
            [
                [
                    row.get("regime", "-"),
                    str(row.get("observations", "-")),
                    _pct(row.get("annual_return")),
                    _num(row.get("sharpe")),
                    _pct(row.get("max_drawdown")),
                ]
                for row in regime_rows
            ],
        )

    confidence = research.get("policy_high_vs_low_confidence", {})
    if confidence:
        document.add_heading("8. Confidence Split", level=1)
        _table(
            document,
            ["Bucket", "CAGR", "Sharpe", "MaxDD"],
            [
                ["High confidence", _pct(confidence["high"]["annual_return"]), _num(confidence["high"]["sharpe"]), _pct(confidence["high"]["max_drawdown"])],
                ["Low confidence", _pct(confidence["low"]["annual_return"]), _num(confidence["low"]["sharpe"]), _pct(confidence["low"]["max_drawdown"])],
            ],
        )
        document.add_paragraph("High-confidence decisions still behave materially better than low-confidence ones, which keeps the policy layer useful even when it is not yet the best benchmark overall.")


def _live_section(document: Document, dashboard: dict) -> None:
    overview = dashboard["overview"]
    document.add_heading("9. Current Live View", level=1)
    _bullet(document, f"Regime: {overview.get('regime', '-')}")
    _bullet(document, f"Beta target: {_pct(overview.get('beta_target'))}")
    _bullet(document, f"Selected hedge: {overview.get('selected_hedge', '-')}")
    _bullet(document, f"Confidence: {_pct(overview.get('confidence'))}")
    _bullet(document, f"Consensus fragility: {_pct(overview.get('consensus_fragility_score'))}")
    _bullet(document, f"Belief-capacity mismatch: {_pct(overview.get('belief_capacity_misalignment'))}")
    for line in overview.get("why_this_action", []):
        _bullet(document, f"Why: {line}")
    for line in overview.get("scenario_narrative", []):
        _bullet(document, f"Scenario: {line}")


def _discovery_section(document: Document, dashboard: dict) -> None:
    document.add_heading("10. Discovery, Holdings Context, And Economic Quality", level=1)
    screener = dashboard.get("screener", {})
    statement = dashboard.get("statement_intelligence", {})
    document.add_paragraph(
        "Discovery is now explicitly separated from current holdings. The screener panel uses `discovery_screener.csv` by default, "
        "while existing positions remain visible in portfolio context."
    )
    top_rows = screener.get("rows", [])[:8]
    if top_rows:
        _table(
            document,
            ["Ticker", "Origin", "Discovery", "Owner elasticity", "Bucket", "Position"],
            [
                [
                    row.get("ticker", "-"),
                    row.get("screen_origin", "-"),
                    _num(row.get("discovery_score")),
                    _pct(row.get("owner_elasticity_score")),
                    row.get("owner_elasticity_bucket", "-") or "-",
                    _pct(row.get("suggested_position")),
                ]
                for row in top_rows
            ],
        )
    top_owner = screener.get("owner_elasticity_top_names", [])[:8]
    if top_owner:
        document.add_paragraph("Top owner-elasticity names:")
        _table(
            document,
            ["Ticker", "Sector", "Elasticity", "Bucket"],
            [
                [
                    row.get("ticker", "-"),
                    row.get("sector", "-") or "-",
                    _pct(row.get("owner_elasticity_score")),
                    row.get("owner_elasticity_bucket", "-") or "-",
                ]
                for row in top_owner
            ],
        )
    top_kernel = statement.get("top_kernel_names", [])[:6]
    if top_kernel:
        document.add_paragraph("Names where accounting strength is best confirmed by cash generation:")
        _table(
            document,
            ["Ticker", "Kernel score", "Cash bucket", "Conviction"],
            [
                [
                    row.get("ticker", "-"),
                    _pct(row.get("earnings_cash_kernel_score")),
                    row.get("earnings_cash_kernel_bucket", "-") or "-",
                    _pct(row.get("statement_conviction_score")),
                ]
                for row in top_kernel
            ],
        )


def _deployment_section(document: Document) -> None:
    document.add_heading("11. Repo Structure And Deployment", level=1)
    _table(
        document,
        ["Area", "Target", "Notes"],
        [
            ["Backend", "Railway", "Python API and local dashboard server, now with `PORT` support and CORS"],
            ["Frontend", "Vercel", "Static app from `src/meta_alpha_allocator/dashboard/static` using configurable `API_BASE`"],
            ["Secrets", "Environment only", "FMP and FRED keys are read from env and are not committed"],
            ["Generated artifacts", "Local only", "`output/`, `cache/`, and `tmp/` are git-ignored"],
        ],
    )
    _bullet(document, "Use `.env.example` as the template for local and deployment configuration.")
    _bullet(document, "Use `railway.toml` or `Procfile` for Railway start-up.")
    _bullet(document, "Use `config.js` to point the Vercel frontend at the Railway API base URL.")


def _limitations(document: Document) -> None:
    document.add_heading("12. Honest Limitations", level=1)
    _bullet(document, "The policy overlay still does not beat the strongest trend-following baseline.")
    _bullet(document, "Consensus fragility and belief-capacity mismatch are conceptually strong but still early as validated live edges.")
    _bullet(document, "Owner elasticity is still a proxy layer built mostly from liquidity, size, and quality rather than direct ownership data.")
    _bullet(document, "The earnings-to-cash kernel is still informational and not allocator core because its out-of-sample utility is mixed.")
    _bullet(document, "A Railway deployment will still require mounted or mirrored access to your local data roots unless those pipelines are moved into cloud storage or scheduled jobs.")


def _glossary(document: Document) -> None:
    document.add_heading("13. Glossary", level=1)
    _bullet(document, "Beta target: the share of the overlay allocated to broad equity risk.")
    _bullet(document, "Hedge: the asset used to offset or soften equity damage.")
    _bullet(document, "Consensus fragility: a measure of how brittle the current market narrative appears.")
    _bullet(document, "Belief-capacity mismatch: a measure of whether capital is too crowded into a narrow narrative relative to its capacity.")
    _bullet(document, "Owner elasticity: a proxy for where smaller proprietary capital may matter more than institutional scale.")
    _bullet(document, "Out-of-sample: periods not used to fit the model at the time the decision was made.")


def build_report() -> Path:
    DOC_ROOT.mkdir(parents=True, exist_ok=True)
    TMP_ROOT.mkdir(parents=True, exist_ok=True)

    research_summary = _load_json(RESEARCH_ROOT / "research_summary.json")
    policy_summary = _load_json(POLICY_ROOT / "policy_backtest_summary.json")
    dashboard = _load_json(DASHBOARD_ROOT / "dashboard_snapshot.json")
    daily_returns = pd.read_csv(RESEARCH_ROOT / "daily_returns.csv", index_col=0, parse_dates=[0]).sort_index()
    policy_returns = pd.read_csv(POLICY_ROOT / "policy_daily_returns.csv", index_col=0, parse_dates=[0]).sort_index()

    growth_frame = pd.DataFrame(index=policy_returns.index)
    if "spy" in daily_returns.columns:
        growth_frame["benchmark_spy"] = (1.0 + daily_returns["spy"].fillna(0.0)).cumprod()
    if "state_overlay" in daily_returns.columns:
        growth_frame["state_overlay"] = (1.0 + daily_returns["state_overlay"].fillna(0.0)).cumprod()
    if "policy_overlay" in policy_returns.columns:
        growth_frame["policy_overlay"] = (1.0 + policy_returns["policy_overlay"].fillna(0.0)).cumprod()
    if "trend_following" in policy_returns.columns:
        growth_frame["trend_following"] = (1.0 + policy_returns["trend_following"].fillna(0.0)).cumprod()
    drawdown_frame = growth_frame.divide(growth_frame.cummax()).subtract(1.0)

    growth_path = TMP_ROOT / "meta_allocator_growth.png"
    drawdown_path = TMP_ROOT / "meta_allocator_drawdown.png"
    _plot_growth(growth_frame, growth_path)
    _plot_drawdown(drawdown_frame, drawdown_path)

    document = Document()
    _style(document)
    _title_block(document, dashboard)
    _executive_summary(document, research_summary, policy_summary, dashboard)
    _plain_language(document, dashboard)
    _system_methodology(document)
    _data_section(document)
    _results_section(document, research_summary, policy_summary, dashboard, growth_path, drawdown_path)
    _live_section(document, dashboard)
    _discovery_section(document, dashboard)
    _deployment_section(document)
    _limitations(document)
    document.add_section(WD_SECTION_START.NEW_PAGE)
    _glossary(document)

    output_path = DOC_ROOT / "meta_allocator_methodology_report.docx"
    document.save(output_path)
    return output_path


if __name__ == "__main__":
    print(build_report())
